import re
import logging
from pathlib import Path
from typing import List, Optional, Tuple, BinaryIO
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""
    text: str
    chunk_index: int
    start_char: int
    end_char: int
    source: str = "text"


@dataclass
class DocumentInfo:
    """Information about a processed document."""
    filename: str
    file_type: str
    total_chars: int
    total_chunks: int
    pages: Optional[int] = None


class DocumentProcessor:
    """
    Service for processing documents and extracting text.
    Supports: PDF, DOCX, DOC, TXT, and plain text.
    """
    
    DEFAULT_CHUNK_SIZE = 2000
    DEFAULT_OVERLAP = 200
    MAX_CHUNK_SIZE = 4000
    
    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        overlap: int = DEFAULT_OVERLAP
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size for each text chunk
            overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = min(chunk_size, self.MAX_CHUNK_SIZE)
        self.overlap = overlap
        
    def extract_text_from_file(
        self,
        file: BinaryIO,
        filename: str
    ) -> Tuple[str, DocumentInfo]:
        """
        Extract text from a file based on its type.
        
        Args:
            file: File-like object
            filename: Original filename (used to determine type)
            
        Returns:
            Tuple of (extracted text, document info)
        """
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            text, pages = self._extract_from_pdf(file)
        elif extension in ['.docx', '.doc']:
            text, pages = self._extract_from_docx(file)
        elif extension == '.txt':
            text = file.read().decode('utf-8', errors='ignore')
            pages = None
        else:
            try:
                text = file.read().decode('utf-8', errors='ignore')
                pages = None
            except Exception:
                raise ValueError(f"Unsupported file format: {extension}")
        
        text = self._clean_text(text)
        
        doc_info = DocumentInfo(
            filename=filename,
            file_type=extension,
            total_chars=len(text),
            total_chunks=0,
            pages=pages
        )
        
        return text, doc_info
    
    def _extract_from_pdf(self, file: BinaryIO) -> Tuple[str, int]:
        """Extract text from PDF file."""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(file)
            pages = len(reader.pages)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n\n'.join(text_parts), pages
            
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF processing. "
                "Install with: pip install pypdf"
            )
    
    def _extract_from_docx(self, file: BinaryIO) -> Tuple[str, int]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts), None
            
        except ImportError:
            raise ImportError(
                "python-docx is required for Word document processing. "
                "Install with: pip install python-docx"
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        text = text.strip()
        return text
    
    def split_into_chunks(
        self,
        text: str,
        source: str = "text"
    ) -> List[TextChunk]:
        """
        Split text into overlapping chunks suitable for LLM processing.
        
        Tries to split at sentence boundaries when possible.
        
        Args:
            text: The text to split
            source: Source identifier for the chunks
            
        Returns:
            List of TextChunk objects
        """
        if not text:
            return []
        
        if len(text) <= self.chunk_size:
            return [TextChunk(
                text=text,
                chunk_index=0,
                start_char=0,
                end_char=len(text),
                source=source
            )]
        
        chunks = []
        current_pos = 0
        chunk_index = 0
        
        while current_pos < len(text):
            end_pos = min(current_pos + self.chunk_size, len(text))
            
            if end_pos < len(text):
                search_start = max(current_pos + self.chunk_size - 200, current_pos)
                search_text = text[search_start:end_pos]
                
                last_boundary = -1
                for match in re.finditer(r'[.!?]\s', search_text):
                    last_boundary = match.end()
                
                if last_boundary > 0:
                    end_pos = search_start + last_boundary
                else:
                    search_text = text[end_pos-50:end_pos]
                    space_pos = search_text.rfind(' ')
                    if space_pos > 0:
                        end_pos = end_pos - 50 + space_pos + 1
            
            chunk_text = text[current_pos:end_pos].strip()
            
            if chunk_text:
                chunks.append(TextChunk(
                    text=chunk_text,
                    chunk_index=chunk_index,
                    start_char=current_pos,
                    end_char=end_pos,
                    source=source
                ))
                chunk_index += 1
            
            if end_pos >= len(text):
                break
            current_pos = max(current_pos + 1, end_pos - self.overlap)
        
        return chunks
    
    def process_document(
        self,
        file: BinaryIO,
        filename: str
    ) -> Tuple[List[TextChunk], DocumentInfo]:
        """
        Full document processing pipeline: extract text and split into chunks.
        
        Args:
            file: File-like object
            filename: Original filename
            
        Returns:
            Tuple of (list of chunks, document info)
        """
        text, doc_info = self.extract_text_from_file(file, filename)
        
        chunks = self.split_into_chunks(text, source=filename)
        
        doc_info.total_chunks = len(chunks)
        
        logger.info(
            f"Processed document '{filename}': "
            f"{doc_info.total_chars} chars, {doc_info.total_chunks} chunks"
        )
        
        return chunks, doc_info
    
    def process_text(self, text: str) -> List[TextChunk]:
        """
        Process plain text input and split into chunks if needed.
        
        Args:
            text: Plain text input
            
        Returns:
            List of TextChunk objects
        """
        text = self._clean_text(text)
        return self.split_into_chunks(text, source="input")


_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the global document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
